from __future__ import annotations

import argparse
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parents[1]
PROJECT_ROOT = REPO_ROOT.parent
REPORT_ASSETS = PROJECT_ROOT / "report_assets"
REPORT_OUTPUTS = PROJECT_ROOT / "report_outputs"
COLAB_RESULTS = PROJECT_ROOT / "colab_results"
GROUPED_SUMMARY = PROJECT_ROOT / "metric_summary_mean_std.csv"


CHECKPOINTS = [
    {
        "file": "checkpoint_01_dataset_setup.docx",
        "title": "Checkpoint 1: Dataset Setup",
        "objective": "Verify the paired underwater dataset and define the supervised restoration task.",
        "key_message": "The active protocol uses grouped source assignments and exact-pair deduplication for more defensible training and testing.",
        "completed": [
            "Verified 3672 matched clear/turbid image pairs.",
            "Created contiguous source-group split files and a duplicate audit.",
            "Separated paired training data from unpaired qualitative data.",
        ],
        "verification": [
            "Checked filename correspondence, readable image dimensions, and one-to-one source-target pairing before training.",
            "Removed 32 exact duplicate records before assigning whole content groups to a single split.",
            "Saved deterministic manifests so every later model uses the same 2,899/357/384 partition.",
        ],
        "decision": "Freeze the grouped manifests; use validation for model selection and reserve the test set for final reporting.",
        "table": [["Split", "Pairs"], ["Train", "2899"], ["Validation", "357"], ["Test", "384"]],
        "figures": [
            (REPORT_ASSETS / "paired_underwater_patches.png", "Example verified paired image patch."),
            (REPORT_ASSETS / "candidate_pair_sheet.png", "Dataset pairing examples used during inspection."),
        ],
        "notes": ["Quantitative metrics are valid only on paired clear/turbid data."],
    },
    {
        "file": "checkpoint_02_diffusion_pipeline.docx",
        "title": "Checkpoint 2: DataLoader and Diffusion Pipeline",
        "objective": "Show that paired images enter the model correctly and that the diffusion training task is implemented.",
        "key_message": "The data pipeline, noising process, and tests are ready before full model training.",
        "completed": [
            "Implemented normalized paired image loading.",
            "Implemented timestep sampling and linear beta schedule.",
            "Added sanity tests for dataset pairing, diffusion, and training steps.",
        ],
        "verification": [
            "Applied synchronized transforms to preserve pixel alignment between underwater inputs and clean references.",
            "Validated timestep ranges, noise tensor shapes, and reverse-sampling output bounds.",
            "Used smoke training and unit tests to detect data-model incompatibilities before full experiments.",
        ],
        "decision": "Keep preprocessing, diffusion schedule, training budget, and callback logic fixed for all compared backbones.",
        "table": [["Component", "Status"], ["Paired DataLoader", "Completed"], ["Forward diffusion", "Completed"], ["Unit tests", "Passing"]],
        "figures": [(REPORT_ASSETS / "proposed_pipeline.png", "Conditional diffusion restoration pipeline.")],
        "notes": ["The same diffusion objective is used for both baseline and residual models."],
    },
    {
        "file": "checkpoint_03_baseline_unet.docx",
        "title": "Checkpoint 3: Baseline U-Net Diffusion Model",
        "objective": "Train and evaluate the baseline conditional U-Net denoising backbone.",
        "key_message": "The baseline model gives the reference point for all later comparisons.",
        "completed": [
            "Implemented conditional U-Net with timestep conditioning.",
            "Completed baseline training and evaluation on the fixed test split.",
            "Generated metric CSV files and qualitative comparison grids.",
        ],
        "verification": [
            "Selected checkpoints from validation behavior without tuning against the held-out test metrics.",
            "Evaluated clipped predictions with shared MSE, MAE, PSNR, SSIM, and Delta E implementations.",
            "Retained per-image outputs and a fixed qualitative grid for traceability beyond aggregate means.",
        ],
        "decision": "Use the default U-Net as the reference architecture and report all later gains relative to this frozen baseline.",
        "table": [["Metric", "Baseline mean +/- SD"], ["MSE", "0.039873 +/- 0.001756"], ["PSNR", "15.780514 +/- 0.187075"], ["SSIM", "0.657180 +/- 0.015721"], ["Delta E", "28.380208 +/- 1.320879"]],
        "figures": [(COLAB_RESULTS / "baseline_full" / "baseline_comparison_grid.png", "Baseline restoration examples from the fixed test split.")],
        "notes": ["This model is the default U-Net diffusion baseline, not the proposed contribution."],
    },
    {
        "file": "checkpoint_04_residual_backbone.docx",
        "title": "Checkpoint 4: Residual Denoising Backbone",
        "objective": "Replace the U-Net denoising backbone with a ResNet-style residual backbone under the same diffusion task.",
        "key_message": "The proposed residual backbone improves structural and color-oriented metrics over the default baseline.",
        "completed": [
            "Implemented residual blocks and residual denoising backbone.",
            "Trained the residual model using the same paired split and objective.",
            "Compared residual and baseline outputs across three seeds on the same 384 test images.",
        ],
        "verification": [
            "Reused the same manifests, optimizer family, diffusion objective, evaluation code, and test tensors.",
            "Stored each seed as an independent run before calculating model-level mean and sample standard deviation.",
            "Checked structural and color metrics separately from pixel-error metrics to avoid a one-number conclusion.",
        ],
        "decision": "Advance the residual backbone to capacity-controlled comparison while qualifying topology and parameter-count differences.",
        "table": [["Metric", "U-Net mean +/- SD", "Residual mean +/- SD"], ["MSE", "0.039873 +/- 0.001756", "0.038782 +/- 0.003119"], ["PSNR", "15.780514 +/- 0.187075", "16.006089 +/- 0.267086"], ["SSIM", "0.657180 +/- 0.015721", "0.779948 +/- 0.006607"], ["Delta E", "28.380208 +/- 1.320879", "27.053617 +/- 1.081249"]],
        "figures": [(REPORT_OUTPUTS / "compact_default_comparison.png", "Matching default U-Net and residual restoration examples.")],
        "notes": ["Residual has better mean SSIM and Delta E than the default U-Net; topology, depth, and capacity differ together."],
    },
    {
        "file": "checkpoint_05_full_results.docx",
        "title": "Checkpoint 5: Full Results",
        "objective": "Evaluate the three main models on the grouped 384-image test split across three seeds.",
        "key_message": "Every final quantitative claim uses the same grouped test list and reports mean +/- sample SD.",
        "completed": [
            "Evaluated 128x128 50-epoch baseline and residual models.",
            "Evaluated default, parameter-matched, and residual models across three seeds.",
            "Generated grouped aggregate metric summary.",
        ],
        "verification": [
            "Ran all three architectures with seeds 42, 123, and 2026 on the identical grouped test list.",
            "Aggregated seed-level scores as mean +/- sample standard deviation rather than pooling test images across runs.",
            "Cross-checked tables, plots, and qualitative examples against the machine-readable grouped summary.",
        ],
        "decision": "Use only the grouped three-model by three-seed matrix for final quantitative claims and keep entropy descriptive.",
        "table": [["Experiment", "PSNR mean +/- SD", "SSIM mean +/- SD", "Delta E mean +/- SD"], ["U-Net", "15.780514 +/- 0.187075", "0.657180 +/- 0.015721", "28.380208 +/- 1.320879"], ["Matched U-Net", "16.017719 +/- 0.277610", "0.722221 +/- 0.022096", "27.590478 +/- 2.280969"], ["Residual", "16.006089 +/- 0.267086", "0.779948 +/- 0.006607", "27.053617 +/- 1.081249"]],
        "figures": [(REPORT_OUTPUTS / "training_loss_curves.png", "Training and validation loss curves."), (REPORT_OUTPUTS / "delta_e_cie76_comparison.png", "Delta E comparison across experiments.")],
        "notes": ["Earlier duration and resolution ablations used the superseded split and are excluded from final claims."],
    },
    {
        "file": "checkpoint_06_ablation_generalization.docx",
        "title": "Checkpoint 6: Parameter-Matched Capacity Control",
        "objective": "Check whether the main architecture comparison is explained by parameter count.",
        "key_message": "The capacity control makes the final claim honest: matched U-Net leads pixel metrics while residual leads mean SSIM and Delta E.",
        "completed": [
            "Completed parameter-matched U-Net capacity control with base_channels=42.",
            "Compared parameter-matched U-Net with the residual model.",
            "Qualified the final conclusion using metric-specific capacity-control evidence.",
        ],
        "verification": [
            "Matched trainable capacity while retaining the plain U-Net topology to isolate architecture from model size.",
            "Confirmed metric direction explicitly: lower for MSE, MAE, and Delta E; higher for PSNR and SSIM.",
            "Recorded mixed leadership and seed variability instead of selecting a universal winner.",
        ],
        "decision": "Conclude that matched U-Net leads mean pixel metrics while the residual model leads mean SSIM and Delta E.",
        "table": [["Metric", "Matched U-Net", "Residual", "Mean leader"], ["MSE", "0.038056 +/- 0.002727", "0.038782 +/- 0.003119", "Matched U-Net"], ["PSNR", "16.017719 +/- 0.277610", "16.006089 +/- 0.267086", "Matched U-Net"], ["SSIM", "0.722221 +/- 0.022096", "0.779948 +/- 0.006607", "Residual"], ["Delta E", "27.590478 +/- 2.280969", "27.053617 +/- 1.081249", "Residual"], ["Entropy", "4.946796 +/- 0.060048", "4.565555 +/- 0.202586", "Descriptive only"]],
        "figures": [(REPORT_OUTPUTS / "compact_capacity_comparison.png", "Matching parameter-matched U-Net and residual restoration examples.")],
        "notes": ["The residual model leads on SSIM and Delta E; the matched U-Net leads on MSE, MAE, and PSNR."],
    },
]


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Create richer supervisor checkpoint DOCX files.")
    parser.add_argument("--output-dir", type=Path, default=Path("docs/supervisor_checkpoints"))
    return parser.parse_args()


def main() -> None:
    if not GROUPED_SUMMARY.exists():
        raise FileNotFoundError(
            "Checkpoint result documents are frozen historical artifacts until the grouped 3-model x 3-seed matrix completes. "
            f"Expected: {GROUPED_SUMMARY}"
        )
    args = parse_args()
    args.output_dir.mkdir(parents=True, exist_ok=True)
    for checkpoint in CHECKPOINTS:
        write_docx(checkpoint, args.output_dir / checkpoint["file"])
    print(f"Wrote {len(CHECKPOINTS)} checkpoint documents to {args.output_dir}")


def write_docx(checkpoint: dict[str, object], output_path: Path) -> None:
    doc = Document()
    setup(doc)
    title(doc, str(checkpoint["title"]))
    section(doc, "Objective")
    para(doc, str(checkpoint["objective"]))
    section(doc, "Key Message")
    callout(doc, str(checkpoint["key_message"]))
    section(doc, "Completed Work")
    for item in checkpoint["completed"]:  # type: ignore[index]
        bullet(doc, str(item))
    section(doc, "Verification and Controls")
    for item in checkpoint["verification"]:  # type: ignore[index]
        bullet(doc, str(item))
    section(doc, "Evidence")
    add_table(doc, checkpoint["table"])  # type: ignore[arg-type]
    section(doc, "Decision / Next Gate")
    callout(doc, str(checkpoint["decision"]))
    section(doc, "Notes")
    for item in checkpoint["notes"]:  # type: ignore[index]
        bullet(doc, str(item))
    for path, text in checkpoint["figures"]:  # type: ignore[index]
        add_picture(doc, Path(path), str(text))
    doc.save(output_path)


def setup(doc: Document) -> None:
    section_obj = doc.sections[0]
    section_obj.start_type = WD_SECTION_START.NEW_PAGE
    section_obj.top_margin = Inches(0.8)
    section_obj.bottom_margin = Inches(0.7)
    section_obj.left_margin = Inches(0.9)
    section_obj.right_margin = Inches(0.9)
    section_obj.footer_distance = Inches(0.3)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    doc.styles["List Bullet"].font.name = "Calibri"
    doc.styles["List Bullet"].font.size = Pt(11)
    footer = section_obj.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_field(footer)


def title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 0, 0)


def section(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(12)


def para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text)


def bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    p.add_run(text)


def callout(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=max(len(row) for row in rows))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            cell = table.cell(r, c)
            cell.text = value
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9)
                    if r == 0:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0, 0, 0)


def add_picture(doc: Document, path: Path, text: str) -> None:
    if not path.exists():
        para(doc, f"Missing figure: {path}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.keep_with_next = True
    with Image.open(path) as image:
        width, height = image.size
    scale = min(5.8 / width, 5.4 / height)
    p.add_run().add_picture(str(path), width=Inches(width * scale), height=Inches(height * scale))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(5)
    run = cap.add_run(text)
    run.italic = True
    run.font.name = "Calibri"
    run.font.size = Pt(8.5)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, end])


if __name__ == "__main__":
    main()
